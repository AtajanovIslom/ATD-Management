"""
Kunlik ish hisobotlari (WorkLog).

Endpoint'lar:
  GET    /api/work-logs                 — mening hisobotlarim (sana filtri bilan)
  POST   /api/work-logs                 — yangi hisobot
  PUT    /api/work-logs/<id>            — tahrirlash (faqat egasi)
  DELETE /api/work-logs/<id>            — o'chirish (faqat egasi)
  GET    /api/work-logs/export          — mening hisobotlarim Word (.docx)

  GET    /api/work-logs/department      — boshqarma rahbari: o'z xodimlari hisobotlari
  GET    /api/work-logs/department/export — o'sha hisobotlar Word (.docx)
"""
import io
from datetime import date as _date, datetime
from flask import Blueprint, request, jsonify, send_file
from flask_jwt_extended import jwt_required, get_jwt, get_jwt_identity
from app import db
from app.models import WorkLog, User
from app.utils import get_scope, is_admin_or_above, log_audit

work_logs_bp = Blueprint('work_logs', __name__)


def _parse_date(s):
    if not s:
        return None
    try:
        return _date.fromisoformat(s[:10])
    except (ValueError, TypeError):
        return None


def _apply_date_range(q, prefix_from='from', prefix_to='to'):
    d_from = _parse_date(request.args.get(prefix_from))
    d_to = _parse_date(request.args.get(prefix_to))
    if d_from:
        q = q.filter(WorkLog.work_date >= d_from)
    if d_to:
        q = q.filter(WorkLog.work_date <= d_to)
    return q, d_from, d_to


# =========================================================================
# XODIMNING O'Z HISOBOTLARI
# =========================================================================

@work_logs_bp.route('', methods=['GET'])
@jwt_required()
def list_mine():
    user_id = int(get_jwt_identity())
    q = WorkLog.query.filter_by(user_id=user_id)
    q, _, _ = _apply_date_range(q)
    items = q.order_by(WorkLog.work_date.desc(), WorkLog.id.desc()).all()
    return jsonify([w.to_dict() for w in items])


@work_logs_bp.route('', methods=['POST'])
@jwt_required()
def create():
    user_id = int(get_jwt_identity())
    data = request.get_json() or {}

    wd = _parse_date(data.get('work_date')) or _date.today()
    content = (data.get('content') or '').strip()
    if not content:
        return jsonify({'error': "Ish tavsifi kiritilishi shart"}), 400

    w = _build_worklog(user_id, wd, content, data)
    db.session.add(w)
    db.session.flush()
    log_audit('create', 'work_log', w.id, entity_label=content[:60])
    db.session.commit()
    return jsonify(w.to_dict()), 201


def _build_worklog(user_id, wd, content, data):
    """Bitta hisobot obyektini yasaydi (bitta yoki batch yaratishda ishlatiladi)."""
    project_id = data.get('project_id') or None
    task_id = data.get('task_id') or None
    interactive_request_id = data.get('interactive_request_id') or None
    return WorkLog(
        user_id=user_id, work_date=wd, content=content,
        project_id=int(project_id) if project_id else None,
        task_id=int(task_id) if task_id else None,
        interactive_request_id=int(interactive_request_id) if interactive_request_id else None,
    )


@work_logs_bp.route('/batch', methods=['POST'])
@jwt_required()
def create_batch():
    """Bir nechta hisobotni bir martada saqlash (jadval ko'rinishida).
       Body: { "items": [ { work_date, content, project_id?, task_id?, interactive_request_id? }, ... ] }
    """
    user_id = int(get_jwt_identity())
    data = request.get_json() or {}
    items = data.get('items') or []
    if not isinstance(items, list) or not items:
        return jsonify({'error': "Kamida bitta hisobot kiritilishi shart"}), 400

    created = []
    for row in items:
        content = (row.get('content') or '').strip()
        if not content:
            continue  # bo'sh qatorlarni o'tkazib yuboramiz
        wd = _parse_date(row.get('work_date')) or _date.today()
        w = _build_worklog(user_id, wd, content, row)
        db.session.add(w)
        created.append(w)

    if not created:
        return jsonify({'error': "Barcha qatorlar bo'sh"}), 400

    db.session.flush()
    log_audit('create', 'work_log', None, entity_label=f"{len(created)} ta kunlik hisobot")
    db.session.commit()
    return jsonify([w.to_dict() for w in created]), 201


@work_logs_bp.route('/<int:log_id>', methods=['PUT'])
@jwt_required()
def update(log_id):
    user_id = int(get_jwt_identity())
    w = WorkLog.query.filter_by(id=log_id, user_id=user_id).first_or_404()
    data = request.get_json() or {}

    if 'work_date' in data:
        wd = _parse_date(data['work_date'])
        if not wd:
            return jsonify({'error': "Sana noto'g'ri"}), 400
        w.work_date = wd
    if 'content' in data:
        content = (data['content'] or '').strip()
        if not content:
            return jsonify({'error': "Ish tavsifi bo'sh bo'lishi mumkin emas"}), 400
        w.content = content
    if 'project_id' in data:
        w.project_id = int(data['project_id']) if data['project_id'] else None
    if 'task_id' in data:
        w.task_id = int(data['task_id']) if data['task_id'] else None
    if 'interactive_request_id' in data:
        w.interactive_request_id = int(data['interactive_request_id']) if data['interactive_request_id'] else None

    log_audit('update', 'work_log', w.id, entity_label=w.content[:60])
    db.session.commit()
    return jsonify(w.to_dict())


@work_logs_bp.route('/<int:log_id>', methods=['DELETE'])
@jwt_required()
def delete(log_id):
    user_id = int(get_jwt_identity())
    w = WorkLog.query.filter_by(id=log_id, user_id=user_id).first_or_404()
    log_audit('delete', 'work_log', w.id, entity_label=w.content[:60])
    db.session.delete(w)
    db.session.commit()
    return jsonify({'message': "O'chirildi"})


# =========================================================================
# WORD (.docx) EKSPORT
# =========================================================================

def _add_logs_table(doc, logs, with_employee=True):
    """Hisobotlar jadvalini hujjatga qo'shadi."""
    from docx.shared import Pt

    cols = ['Sana', 'Xodim', 'Loyiha / Vazifa / Interaktiv', 'Bajarilgan ish'] if with_employee \
        else ['Sana', 'Loyiha / Vazifa / Interaktiv', 'Bajarilgan ish']
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for i, name in enumerate(cols):
        hdr[i].text = name
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    for w in logs:
        row = table.add_row().cells
        if with_employee:
            row[0].text = w.work_date.strftime('%d.%m.%Y') if w.work_date else ''
            row[1].text = w.user.full_name if w.user else ''
            row[2].text = w.ref_label()
            row[3].text = w.content or ''
        else:
            row[0].text = w.work_date.strftime('%d.%m.%Y') if w.work_date else ''
            row[1].text = w.ref_label()
            row[2].text = w.content or ''


def _build_docx(title, subtitle, logs, group_by_division=False):
    """Hisobotlar ro'yxatidan .docx hujjat yasaydi va BytesIO qaytaradi.

    group_by_division=True bo'lsa — hisobotlar bo'limlar kesimida guruhlanadi
    (boshqarma rahbari eksporti uchun).
    """
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    h = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subtitle:
        p = doc.add_paragraph(subtitle)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True

    if not logs:
        doc.add_paragraph("Ushbu davr uchun hisobot topilmadi.")
    elif group_by_division:
        # Bo'limlar kesimida guruhlaymiz
        groups = {}
        for w in logs:
            div_name = (w.user.division.name if w.user and w.user.division else None) or "Bo'lim ko'rsatilmagan"
            groups.setdefault(div_name, []).append(w)
        for div_name in sorted(groups.keys()):
            doc.add_heading(f"🏢 {div_name}", level=2)
            _add_logs_table(doc, groups[div_name], with_employee=True)
            doc.add_paragraph()
    else:
        _add_logs_table(doc, logs, with_employee=True)

    doc.add_paragraph()
    footer = doc.add_paragraph(f"Yaratildi: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    footer.runs[0].font.size = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _range_subtitle(d_from, d_to):
    if d_from and d_to:
        return f"Davr: {d_from.strftime('%d.%m.%Y')} — {d_to.strftime('%d.%m.%Y')}"
    if d_from:
        return f"{d_from.strftime('%d.%m.%Y')} dan"
    if d_to:
        return f"{d_to.strftime('%d.%m.%Y')} gacha"
    return "Barcha davr"


@work_logs_bp.route('/export', methods=['GET'])
@jwt_required()
def export_mine():
    user_id = int(get_jwt_identity())
    user = User.query.get(user_id)
    q = WorkLog.query.filter_by(user_id=user_id)
    q, d_from, d_to = _apply_date_range(q)
    logs = q.order_by(WorkLog.work_date.asc(), WorkLog.id.asc()).all()

    buf = _build_docx(
        title="Kunlik ish hisoboti",
        subtitle=f"{user.full_name if user else ''} | {_range_subtitle(d_from, d_to)}",
        logs=logs,
    )
    fname = f"hisobot_{_date.today().isoformat()}.docx"
    return send_file(buf, as_attachment=True, download_name=fname,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


# =========================================================================
# BOSHQARMA RAHBARI — O'Z XODIMLARI HISOBOTLARI
# =========================================================================

def _department_scope_user_ids(role, dept_id, div_id, self_id):
    """Boshqarma rahbari o'z boshqarmasidagi xodimlar hisobotlarini ko'radi.
       To'liq huquqli rollar — barcha xodimlarni.
    """
    if role in ('superadmin', 'director', 'deputy_director'):
        return None  # hammasi
    if role == 'admin' and dept_id:
        users = User.query.filter_by(department_id=dept_id).all()
        return {u.id for u in users}
    if role == 'department_admin' and div_id:
        users = User.query.filter_by(division_id=div_id).all()
        return {u.id for u in users}
    return set()


def _department_query(role, dept_id, div_id, self_id):
    q = WorkLog.query
    ids = _department_scope_user_ids(role, dept_id, div_id, self_id)
    if ids is not None:
        if not ids:
            return q.filter(db.false())
        q = q.filter(WorkLog.user_id.in_(ids))
    # Ixtiyoriy: bitta xodim bo'yicha filtr
    uid = request.args.get('user_id', type=int)
    if uid:
        q = q.filter(WorkLog.user_id == uid)
    return q


@work_logs_bp.route('/department', methods=['GET'])
@jwt_required()
def list_department():
    role, dept_id, div_id = get_scope(get_jwt())
    self_id = int(get_jwt_identity())
    if not is_admin_or_above(role) and role != 'department_admin':
        return jsonify({'error': "Ruxsat yo'q"}), 403

    q = _department_query(role, dept_id, div_id, self_id)
    q, _, _ = _apply_date_range(q)
    items = q.order_by(WorkLog.work_date.desc(), WorkLog.id.desc()).all()
    return jsonify([w.to_dict() for w in items])


@work_logs_bp.route('/department/export', methods=['GET'])
@jwt_required()
def export_department():
    role, dept_id, div_id = get_scope(get_jwt())
    self_id = int(get_jwt_identity())
    if not is_admin_or_above(role) and role != 'department_admin':
        return jsonify({'error': "Ruxsat yo'q"}), 403

    actor = User.query.get(self_id)
    q = _department_query(role, dept_id, div_id, self_id)
    q, d_from, d_to = _apply_date_range(q)
    logs = q.order_by(WorkLog.work_date.asc(), WorkLog.user_id.asc(), WorkLog.id.asc()).all()

    buf = _build_docx(
        title="Boshqarma xodimlari kunlik hisobotlari",
        subtitle=_range_subtitle(d_from, d_to),
        logs=logs,
        group_by_division=True,  # bo'limlar kesimida
    )
    fname = f"boshqarma_hisobot_{_date.today().isoformat()}.docx"
    return send_file(buf, as_attachment=True, download_name=fname,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
